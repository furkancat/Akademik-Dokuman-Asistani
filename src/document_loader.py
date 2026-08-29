"""
Yerel Akademik Doküman Asistanı - Doküman Yükleme Modülü

Desteklenen formatlar: PDF, TXT, DOCX
"""

import logging
from pathlib import Path
from typing import List, Optional

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredWordDocumentLoader,
)
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import CHUNK_SIZE, CHUNK_OVERLAP, DOCUMENTS_DIR

logger = logging.getLogger(__name__)


class DocumentLoader:
    """PDF, TXT, DOCX dosyalarını yükleyip chunk'lara ayıran sınıf."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".doc", ".md"}

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        logger.info(f"DocumentLoader başlatıldı: chunk_size={chunk_size}, overlap={chunk_overlap}")

    def load_pdf(self, file_path: Path) -> List[Document]:
        """PDF dosyasını yükle."""
        logger.info(f"PDF yükleniyor: {file_path}")
        try:
            loader = PyPDFLoader(str(file_path))
            documents = loader.load()
            # Metadata'ya kaynak ekle
            for doc in documents:
                doc.metadata["source"] = str(file_path)
                doc.metadata["file_type"] = "pdf"
                doc.metadata["filename"] = file_path.name
            logger.info(f"PDF yüklendi: {len(documents)} sayfa")
            return documents
        except Exception as e:
            logger.error(f"PDF yükleme hatası ({file_path}): {e}")
            return []

    def load_txt(self, file_path: Path) -> List[Document]:
        """TXT dosyasını yükle."""
        logger.info(f"TXT yükleniyor: {file_path}")
        try:
            loader = TextLoader(str(file_path), encoding="utf-8")
            documents = loader.load()
            for doc in documents:
                doc.metadata["source"] = str(file_path)
                doc.metadata["file_type"] = "txt"
                doc.metadata["filename"] = file_path.name
            logger.info(f"TXT yüklendi: {len(documents)} belge")
            return documents
        except UnicodeDecodeError:
            # UTF-8 başarısız olursa diğer encoding'leri dene
            for encoding in ["latin-1", "cp1254", "iso-8859-9"]:
                try:
                    loader = TextLoader(str(file_path), encoding=encoding)
                    documents = loader.load()
                    for doc in documents:
                        doc.metadata["source"] = str(file_path)
                        doc.metadata["file_type"] = "txt"
                        doc.metadata["filename"] = file_path.name
                    logger.info(f"TXT yüklendi (encoding={encoding}): {len(documents)} belge")
                    return documents
                except Exception:
                    continue
            logger.error(f"TXT yükleme hatası - hiçbir encoding çalışmadı: {file_path}")
            return []
        except Exception as e:
            logger.error(f"TXT yükleme hatası ({file_path}): {e}")
            return []

    def load_docx(self, file_path: Path) -> List[Document]:
        """DOCX dosyasını yükle."""
        logger.info(f"DOCX yükleniyor: {file_path}")
        try:
            loader = UnstructuredWordDocumentLoader(str(file_path))
            documents = loader.load()
            for doc in documents:
                doc.metadata["source"] = str(file_path)
                doc.metadata["file_type"] = "docx"
                doc.metadata["filename"] = file_path.name
            logger.info(f"DOCX yüklendi: {len(documents)} belge")
            return documents
        except Exception as e:
            logger.error(f"DOCX yükleme hatası ({file_path}): {e}")
            # Alternatif: python-docx ile dene
            try:
                return self._load_docx_alternative(file_path)
            except Exception as e2:
                logger.error(f"Alternatif DOCX yükleme de başarısız: {e2}")
                return []

    def _load_docx_alternative(self, file_path: Path) -> List[Document]:
        """python-docx ile DOCX yükleme (yedek yöntem)."""
        try:
            import docx
            doc = docx.Document(str(file_path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            content = "\n".join(full_text)
            document = Document(
                page_content=content,
                metadata={
                    "source": str(file_path),
                    "file_type": "docx",
                    "filename": file_path.name,
                }
            )
            logger.info(f"DOCX (alternatif) yüklendi: {file_path.name}")
            return [document]
        except ImportError:
            logger.error("python-docx kurulu değil. Yüklemek için: pip install python-docx")
            return []

    def load_single_file(self, file_path: Path) -> List[Document]:
        """Tek bir dosyayı formatına göre yükle."""
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"Dosya bulunamadı: {file_path}")
            return []

        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self.load_pdf(file_path)
        elif extension in [".txt", ".md"]:
            return self.load_txt(file_path)
        elif extension in [".docx", ".doc"]:
            return self.load_docx(file_path)
        else:
            logger.warning(f"Desteklenmeyen dosya formatı: {extension}")
            return []

    def load_directory(
        self,
        directory: Path = DOCUMENTS_DIR,
        recursive: bool = True,
    ) -> List[Document]:
        """
        Bir dizindeki tüm desteklenen dosyaları yükle.

        Args:
            directory: Tarancak dizin
            recursive: Alt dizinleri de tara

        Returns:
            List[Document]: Yüklenen tüm dokümanlar
        """
        directory = Path(directory)
        if not directory.exists():
            logger.warning(f"Dizin bulunamadı: {directory}")
            return []

        all_documents = []

        if recursive:
            files = [
                f for f in directory.rglob("*")
                if f.suffix.lower() in self.SUPPORTED_EXTENSIONS
            ]
        else:
            files = [
                f for f in directory.iterdir()
                if f.suffix.lower() in self.SUPPORTED_EXTENSIONS
            ]

        logger.info(f"{directory} dizininde {len(files)} dosya bulundu")

        for file_path in files:
            docs = self.load_single_file(file_path)
            all_documents.extend(docs)

        logger.info(f"Toplam {len(all_documents)} ham belge yüklendi")
        return all_documents

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        Dokümanları chunk'lara böl.

        Args:
            documents: Bölünecek doküman listesi

        Returns:
            List[Document]: Chunk'lanmış dokümanlar
        """
        if not documents:
            logger.warning("Bölünecek doküman yok")
            return []

        chunks = self.text_splitter.split_documents(documents)
        logger.info(f"{len(documents)} belge -> {len(chunks)} chunk'a bölündü")
        return chunks

    def load_and_split(
        self,
        source: Path,
        recursive: bool = True,
    ) -> List[Document]:
        """
        Dosya veya dizini yükle ve chunk'lara böl (tek adım).

        Args:
            source: Dosya veya dizin yolu
            recursive: Alt dizinleri de tara (dizin için)

        Returns:
            List[Document]: Chunk'lanmış dokümanlar
        """
        source = Path(source)

        if source.is_file():
            documents = self.load_single_file(source)
        elif source.is_dir():
            documents = self.load_directory(source, recursive=recursive)
        else:
            logger.error(f"Geçersiz kaynak: {source}")
            return []

        return self.split_documents(documents)


# === KOLAY KULLANIM FONKSİYONLARI ===

def load_documents(
    source: Optional[Path] = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> List[Document]:
    """
    Dokümanları yükle ve chunk'lara böl (kolay kullanım).

    Args:
        source: Dosya veya dizin yolu (None = varsayılan documents dizini)
        chunk_size: Her chunk'ın karakter sayısı
        chunk_overlap: Chunk'lar arası örtüşme

    Returns:
        List[Document]: Chunk'lanmış dokümanlar
    """
    loader = DocumentLoader(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    source = source or DOCUMENTS_DIR
    return loader.load_and_split(source)


def add_document(file_path: Path) -> List[Document]:
    """
    Tek bir dosyayı asistanın bellek dizinine kopyala ve yükle.

    Args:
        file_path: Kopyalanacak dosya yolu

    Returns:
        List[Document]: Yüklenen chunk'lar
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")

    # Dosyayı documents dizinine kopyala
    dest_path = DOCUMENTS_DIR / file_path.name
    import shutil
    shutil.copy2(file_path, dest_path)
    logger.info(f"Dosya kopyalandı: {file_path} -> {dest_path}")

    # Yükle ve chunk'la
    loader = DocumentLoader()
    return loader.load_and_split(dest_path)


# === DOĞRUDAN ÇALIŞTIRMA ===
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    # Test
    loader = DocumentLoader()

    # Dizin içeriğini listele
    print(f"\n=== {DOCUMENTS_DIR} dizinindeki dosyalar ===")
    if DOCUMENTS_DIR.exists():
        for f in DOCUMENTS_DIR.rglob("*"):
            if f.suffix.lower() in DocumentLoader.SUPPORTED_EXTENSIONS:
                print(f"  - {f.relative_to(DOCUMENTS_DIR)}")

    # Dokümanları yükle
    chunks = loader.load_and_split(DOCUMENTS_DIR)
    print(f"\nToplam {len(chunks)} chunk oluşturuldu")

    if chunks:
        print(f"\nİlk chunk örneği:")
        print(f"  Kaynak: {chunks[0].metadata.get('source', 'N/A')}")
        print(f"  İçerik (ilk 200 karakter): {chunks[0].page_content[:200]}...")
